"""
News Reader — Google News RSS + BBC RSS (no API key needed), optional
NewsAPI merge, LLM-written spoken briefings, ordinal follow-ups
("tell me more about the third one"), and "save the news" to a .docx.
Results cached for NEWS_CACHE_MINUTES so follow-ups are instant.
"""
import json
import re
import time
from urllib.parse import quote_plus

import feedparser
import requests
from bs4 import BeautifulSoup

from config import Config
from brain.prompts import NEWS_BRIEFING_PROMPT, NEWS_DEEPDIVE_PROMPT

UA = {"User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) JARVIS/1.0"}

FEEDS = {
    "top": [
        "https://news.google.com/rss?hl=en-US&gl=US&ceid=US:en",
        "http://feeds.bbci.co.uk/news/rss.xml",
    ],
    "world": ["http://feeds.bbci.co.uk/news/world/rss.xml"],
    "technology": [
        "http://feeds.bbci.co.uk/news/technology/rss.xml",
        "https://news.google.com/rss/headlines/section/topic/TECHNOLOGY?hl=en-US&gl=US&ceid=US:en",
    ],
    "tech": ["http://feeds.bbci.co.uk/news/technology/rss.xml"],
    "business": ["http://feeds.bbci.co.uk/news/business/rss.xml"],
    "science": ["http://feeds.bbci.co.uk/news/science_and_environment/rss.xml"],
    "health": ["http://feeds.bbci.co.uk/news/health/rss.xml"],
    "sports": ["http://feeds.bbci.co.uk/sport/rss.xml"],
    "sport": ["http://feeds.bbci.co.uk/sport/rss.xml"],
    "entertainment": ["http://feeds.bbci.co.uk/news/entertainment_and_arts/rss.xml"],
    "uae": ["https://news.google.com/rss/search?q=UAE&hl=en-US&gl=US&ceid=US:en"],
}

ORDINALS = {
    "first": 0, "1st": 0, "one": 0, "1": 0,
    "second": 1, "2nd": 1, "two": 1, "2": 1,
    "third": 2, "3rd": 2, "three": 2, "3": 2,
    "fourth": 3, "4th": 3, "four": 3, "4": 3,
    "fifth": 4, "5th": 4, "five": 4, "5": 4,
    "sixth": 5, "6th": 5, "six": 5, "6": 5,
    "seventh": 6, "7th": 6, "seven": 6, "7": 6,
}


# ---------------------------------------------------------------------------
# Fetching
# ---------------------------------------------------------------------------
def _norm_title(t):
    return re.sub(r"\W+", "", (t or "").lower())


def _parse_feed(url, limit=10):
    out = []
    try:
        feed = feedparser.parse(url, request_headers=UA)
        for entry in feed.entries[:limit]:
            title = getattr(entry, "title", "") or ""
            link = getattr(entry, "link", "") or ""
            source = ""
            try:
                source = entry.source.get("title", "")  # type: ignore[attr-defined]
            except Exception:
                pass
            if not source:
                source = getattr(entry, "author", "") or feed.feed.get("title", "")
            published = getattr(entry, "published", "") or getattr(entry, "updated", "")
            out.append({"title": title.strip(), "link": link.strip(),
                        "source": source.strip(), "published": published})
    except Exception:
        pass
    return out


def _newsapi(topic=None, limit=8):
    if not Config.NEWS_API_KEY:
        return []
    try:
        if topic:
            url = ("https://newsapi.org/v2/everything?q=" + quote_plus(topic) +
                   f"&pageSize={limit}&sortBy=publishedAt&apiKey={Config.NEWS_API_KEY}")
        else:
            url = ("https://newsapi.org/v2/top-headlines?language=en&pageSize="
                   f"{limit}&apiKey={Config.NEWS_API_KEY}")
        resp = requests.get(url, timeout=10, headers=UA)
        data = resp.json()
        out = []
        for a in data.get("articles", []):
            out.append({
                "title": (a.get("title") or "").strip(),
                "link": (a.get("url") or "").strip(),
                "source": (a.get("source") or {}).get("name", ""),
                "published": a.get("publishedAt", "") or "",
            })
        return out
    except Exception:
        return []


def fetch_headlines(topic=None, limit=12):
    topic = (topic or "").strip().lower()
    if topic and topic in FEEDS:
        urls = FEEDS[topic]
    elif topic:
        urls = ["https://news.google.com/rss/search?q=" + quote_plus(topic) +
                "&hl=en-US&gl=US&ceid=US:en"]
    else:
        urls = FEEDS["top"]

    items = []
    for u in urls:
        items.extend(_parse_feed(u, limit=limit))
    items.extend(_newsapi(topic or None))

    seen = set()
    deduped = []
    for it in items:
        key = _norm_title(it["title"])[:80]
        if not key or key in seen:
            continue
        seen.add(key)
        deduped.append(it)
        if len(deduped) >= limit:
            break
    return deduped


# ---------------------------------------------------------------------------
# Cache
# ---------------------------------------------------------------------------
def _load_cache():
    try:
        if Config.NEWS_CACHE_FILE.exists():
            data = json.loads(Config.NEWS_CACHE_FILE.read_text(encoding="utf-8"))
            age_min = (time.time() - data.get("time", 0)) / 60.0
            if age_min <= Config.NEWS_CACHE_MINUTES:
                return data
    except Exception:
        pass
    return None


def _save_cache(topic, headlines):
    try:
        Config.NEWS_CACHE_FILE.write_text(json.dumps({
            "time": time.time(), "topic": topic or "",
            "headlines": headlines,
        }, ensure_ascii=False, indent=2), encoding="utf-8")
    except Exception:
        pass


# ---------------------------------------------------------------------------
# Article extraction
# ---------------------------------------------------------------------------
def fetch_article_text(url, max_chars=6000):
    try:
        resp = requests.get(url, timeout=12, headers=UA, allow_redirects=True)
        soup = BeautifulSoup(resp.text, "lxml")
        for tag in soup(["script", "style", "noscript", "header", "footer",
                         "nav", "aside", "form"]):
            tag.decompose()
        paragraphs = [p.get_text(" ", strip=True) for p in soup.find_all("p")]
        text = "\n".join(p for p in paragraphs if len(p) > 40)
        return text[:max_chars]
    except Exception:
        return ""


# ---------------------------------------------------------------------------
# Actions
# ---------------------------------------------------------------------------
def briefing(topic, ctx):
    topic = (topic or "").strip()
    cached = _load_cache()
    if cached and (cached.get("topic") or "").lower() == topic.lower():
        headlines = cached["headlines"]
    else:
        headlines = fetch_headlines(topic or None, limit=12)
        _save_cache(topic, headlines)

    if not headlines:
        return ("I couldn't pull any headlines just now, sir — "
                "the feeds may be unreachable.")

    ctx.state["news"] = {"topic": topic, "headlines": headlines,
                         "time": time.time()}

    top = headlines[:7]
    block = "\n".join(
        f"{i+1}. {h['title']} ({h['source']})" for i, h in enumerate(top)
    )
    script = ctx.llm.quick(
        NEWS_BRIEFING_PROMPT.format(n=len(top), headlines=block),
        max_tokens=500,
    ) if ctx.llm.available else ""
    if not script:
        label = f" on {topic}" if topic else ""
        parts = [f"Here are the top stories{label}, sir."]
        for i, h in enumerate(top, 1):
            parts.append(f"Story {i}: {h['title']}.")
        parts.append("Ask me about any of them for more detail.")
        script = " ".join(parts)
    return script


def more_about(ordinal, ctx):
    news = ctx.state.get("news")
    if not news or not news.get("headlines"):
        cached = _load_cache()
        if cached:
            news = cached
        else:
            return "There's no briefing in memory, sir. Ask for the news first."

    idx = ORDINALS.get(str(ordinal).strip().lower())
    if idx is None:
        m = re.search(r"(\d+)", str(ordinal))
        if m:
            idx = int(m.group(1)) - 1
    headlines = news.get("headlines", [])
    if idx is None or idx < 0 or idx >= len(headlines):
        return "Which story did you mean, sir? Give me a number."

    item = headlines[idx]
    text = fetch_article_text(item["link"])
    if not text:
        return (f"I couldn't fetch the full article, sir, but the headline is: "
                f"{item['title']}, from {item['source']}.")
    summary = ctx.llm.quick(
        NEWS_DEEPDIVE_PROMPT.format(article=text[:4000]),
        max_tokens=300,
    ) if ctx.llm.available else ""
    if not summary:
        summary = text[:500]
    return f"On story {idx + 1}: {summary}"


def save_news(ctx):
    news = ctx.state.get("news") or _load_cache() or {}
    headlines = news.get("headlines") or []
    if not headlines:
        return "There's no briefing to save, sir. Ask for the news first."
    try:
        from docx import Document
        from docx.shared import Pt
        doc = Document()
        doc.add_heading("JARVIS News Briefing", 0)
        doc.add_paragraph(time.strftime("Generated %A, %d %B %Y at %H:%M"))
        topic = news.get("topic")
        if topic:
            doc.add_paragraph(f"Topic: {topic}")
        for i, h in enumerate(headlines[:10], 1):
            doc.add_heading(f"{i}. {h['title']}", level=2)
            p = doc.add_paragraph()
            p.add_run(f"Source: {h['source']}\n").italic = True
            if h.get("published"):
                p.add_run(f"Published: {h['published']}\n").italic = True
            p.add_run(h["link"])
            for run in p.runs:
                run.font.size = Pt(9)
        ts = time.strftime("%Y%m%d_%H%M%S")
        path = Config.DESKTOP_PATH / f"news_briefing_{ts}.docx"
        doc.save(str(path))
        return f"Briefing saved to your desktop as {path.name}, sir."
    except Exception as exc:
        return f"I couldn't save the briefing: {exc}."


# ---------------------------------------------------------------------------
# Skill dispatch entry
# ---------------------------------------------------------------------------
def handle(intent, ctx):
    skill = intent.get("skill")
    params = intent.get("params", {}) or {}
    if skill == "news.latest":
        return briefing("", ctx)
    if skill == "news.topic":
        return briefing(params.get("topic", ""), ctx)
    if skill == "news.more":
        return more_about(params.get("ordinal", ""), ctx)
    if skill == "news.save":
        return save_news(ctx)
    return None
